"""Small production-friendly document parser migrated from enrui-ai-platform."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument
from fastapi import HTTPException
from pypdf import PdfReader


def parse_document(raw: bytes, filename: str) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return _parse_pdf(raw)
    if suffix == ".docx":
        return _parse_docx(raw)
    if suffix in {".txt", ".md", ".csv", ".json"}:
        return _decode_text(raw)
    raise HTTPException(status_code=415, detail="仅支持 PDF、DOCX、TXT、Markdown、CSV、JSON")


def _decode_text(raw: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gbk", "latin-1"):
        try:
            content = raw.decode(encoding).strip()
            if content:
                return content
        except UnicodeDecodeError:
            continue
    raise HTTPException(status_code=422, detail="无法解析文本文件编码")


def _parse_pdf(raw: bytes) -> str:
    try:
        reader = PdfReader(BytesIO(raw))
        if reader.is_encrypted:
            raise HTTPException(status_code=422, detail="PDF 已加密，无法解析")
        pages = [
            f"[第 {index} 页]\n{content}"
            for index, page in enumerate(reader.pages, start=1)
            if (content := (page.extract_text() or "").strip())
        ]
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=422, detail="PDF 文件损坏或无法读取") from exc
    if not pages:
        raise HTTPException(status_code=422, detail="PDF 未提取到文本，扫描件需后续 OCR 服务")
    return "\n\n".join(pages)


def _parse_docx(raw: bytes) -> str:
    try:
        document = DocxDocument(BytesIO(raw))
    except Exception as exc:
        raise HTTPException(status_code=422, detail="DOCX 文件损坏或无法读取") from exc
    blocks = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            blocks.append(" | ".join(cell.text.strip() for cell in row.cells))
    if not blocks:
        raise HTTPException(status_code=422, detail="DOCX 中未找到可解析文本")
    return "\n\n".join(blocks)
